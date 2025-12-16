import os
import json
import streamlit as st
from datetime import datetime
from openai import OpenAI
from openai import RateLimitError, APIError, APITimeoutError

# -------------------- OpenAI client -------------------- #

def get_client():
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        return None
    return OpenAI(api_key=key)

# -------------------- Prompting -------------------- #

DUAL_OUTPUT_SYSTEM_PROMPT = """
You are a senior roofing service manager and technical writer.

You translate vague, shorthand roofer field notes into professional documentation.

You must produce TWO versions of the same report:
1) INTERNAL / MANUFACTURER-SAFE (authoritative default)
2) CUSTOMER-FRIENDLY (translation derived from internal)

Key behavior:
- Roofers often omit details; you may reasonably infer implications using common roofing practice.
- Use manufacturer-recognizable language (proper adhesion, acceptable installation conditions, substrate conditions)
  WITHOUT citing exact specs, temperature numbers, product SKUs, manufacturer names, or warranty conclusions.
- Do NOT assign blame.
- Use qualified language (“likely”, “may”, “appears”) when certainty is limited.

Active leaks:
- Set active_leak_reported = true ONLY if the notes explicitly state an active leak or water intrusion.
- If not stated, set false.
- You may mention “risk of leakage” only as a potential concern if supported by the notes.

Output MUST be valid JSON only (no markdown, no extra text).

Return JSON with EXACTLY these keys:

{
  "internal_report": {
    "service_summary": string,
    "roof_system": string,
    "primary_issue": string,
    "location": string,
    "active_leak_reported": boolean,
    "observations": [string],
    "installation_site_conditions": [string],
    "potential_concerns": [string],
    "recommended_next_steps": [string],
    "severity": "Low"|"Moderate"|"High",
    "urgency": "Routine"|"Soon"|"Immediate"
  },
  "customer_report": {
    "what_we_found": string,
    "why_this_matters": string,
    "what_this_could_lead_to": [string],
    "recommended_next_steps": [string],
    "priority": "Routine"|"Soon"|"Immediate"
  },
  "clarifying_questions": [string]
}

Rules:
- If unknown, use "Not specified" for strings, [] for lists, and false for booleans (unless explicitly true).
- The customer_report must NOT introduce new facts beyond internal_report.
"""

# -------------------- JSON safety & validation -------------------- #

def safe_json_load(text: str) -> dict:
    """
    Strict JSON parse with a small salvage attempt if the model wraps JSON.
    """
    text = (text or "").strip()
    try:
        return json.loads(text)
    except Exception:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(text[start:end+1])
        raise

def _as_str(v, default="Not specified"):
    if v is None:
        return default
    s = str(v).strip()
    return s if s else default

def _as_bool(v):
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        t = v.strip().lower()
        if t in {"true", "yes", "y"}:
            return True
        if t in {"false", "no", "n"}:
            return False
    return False

def _as_list(v):
    if v is None:
        return []
    if isinstance(v, list):
        return [str(x).strip() for x in v if str(x).strip()]
    if isinstance(v, str) and v.strip():
        return [v.strip()]
    return []

SEVERITY_SET = {"Low", "Moderate", "High"}
URGENCY_SET = {"Routine", "Soon", "Immediate"}

def normalize_and_validate_dual(data: dict) -> dict:
    """
    Ensures required structure exists and types are sane.
    Fills missing fields with safe defaults.
    Ignores extra keys safely.
    """
    if not isinstance(data, dict):
        raise ValueError("Parsed JSON is not an object.")

    internal = data.get("internal_report") if isinstance(data.get("internal_report"), dict) else {}
    customer = data.get("customer_report") if isinstance(data.get("customer_report"), dict) else {}

    out = {
        "internal_report": {
            "service_summary": _as_str(internal.get("service_summary")),
            "roof_system": _as_str(internal.get("roof_system")),
            "primary_issue": _as_str(internal.get("primary_issue")),
            "location": _as_str(internal.get("location")),
            "active_leak_reported": _as_bool(internal.get("active_leak_reported")),
            "observations": _as_list(internal.get("observations")),
            "installation_site_conditions": _as_list(internal.get("installation_site_conditions")),
            "potential_concerns": _as_list(internal.get("potential_concerns")),
            "recommended_next_steps": _as_list(internal.get("recommended_next_steps")),
            "severity": _as_str(internal.get("severity"), "Moderate").title(),
            "urgency": _as_str(internal.get("urgency"), "Soon").title(),
        },
        "customer_report": {
            "what_we_found": _as_str(customer.get("what_we_found")),
            "why_this_matters": _as_str(customer.get("why_this_matters")),
            "what_this_could_lead_to": _as_list(customer.get("what_this_could_lead_to")),
            "recommended_next_steps": _as_list(customer.get("recommended_next_steps")),
            "priority": _as_str(customer.get("priority"), "Soon").title(),
        },
        "clarifying_questions": _as_list(data.get("clarifying_questions")),
    }

    if out["internal_report"]["severity"] not in SEVERITY_SET:
        out["internal_report"]["severity"] = "Moderate"
    if out["internal_report"]["urgency"] not in URGENCY_SET:
        out["internal_report"]["urgency"] = "Soon"
    if out["customer_report"]["priority"] not in URGENCY_SET:
        out["customer_report"]["priority"] = out["internal_report"]["urgency"]

    return out

# -------------------- OpenAI call -------------------- #

def normalize_notes_with_gpt(client: OpenAI, notes: str):
    """
    Returns: (data_or_None, raw_json_or_None, err_message_or_None)
    Never throws to Streamlit UI.
    """
    try:
        completion = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": DUAL_OUTPUT_SYSTEM_PROMPT},
                {"role": "user", "content": notes},
            ],
            response_format={"type": "json_object"},  # ✅ JSON mode
            temperature=0.2,
            max_tokens=1100,
        )
        raw = completion.choices[0].message.content
        parsed = safe_json_load(raw)
        validated = normalize_and_validate_dual(parsed)
        return validated, raw, None

    except RateLimitError:
        return None, None, (
            "AI is temporarily unavailable due to usage/rate limits (often billing/credits/limits). "
            "Use Manual Normalize below or try again later."
        )
    except (APITimeoutError, APIError) as e:
        return None, None, f"AI service error: {e}"
    except json.JSONDecodeError:
        return None, None, (
            "AI returned a response that was not valid JSON. "
            "Try again, or use Manual Normalize below."
        )
    except Exception as e:
        return None, None, f"Unexpected error: {e}"

# -------------------- Report rendering -------------------- #

def _bullets(items):
    if not items:
        return "- Not specified"
    return "\n".join([f"- {x}" for x in items])

def build_internal_report(internal: dict) -> str:
    leak = "Yes" if internal.get("active_leak_reported") else "No / Not reported"

    return f"""
## INTERNAL / MANUFACTURER-SAFE REPORT

**Service Summary**  
{internal.get("service_summary", "Not specified")}

**Roof System**  
{internal.get("roof_system", "Not specified")}

**Primary Issue**  
{internal.get("primary_issue", "Not specified")}

**Location**  
{internal.get("location", "Not specified")}

**Active Leak Reported**  
**{leak}**

**Observations**
{_bullets(internal.get("observations", []))}

**Installation / Site Conditions**
{_bullets(internal.get("installation_site_conditions", []))}

**Potential Concerns**
{_bullets(internal.get("potential_concerns", []))}

**Recommended Next Steps**
{_bullets(internal.get("recommended_next_steps", []))}

**Severity / Urgency**
- **Severity:** {internal.get("severity", "Moderate")}
- **Urgency:** {internal.get("urgency", "Soon")}
""".strip()

def build_customer_report(customer: dict) -> str:
    return f"""
## CUSTOMER-FRIENDLY REPORT

**What We Found**  
{customer.get("what_we_found", "Not specified")}

**Why This Matters**  
{customer.get("why_this_matters", "Not specified")}

**What This Could Lead To**
{_bullets(customer.get("what_this_could_lead_to", []))}

**Recommended Next Steps**
{_bullets(customer.get("recommended_next_steps", []))}

**Priority**
**{customer.get("priority", "Soon")}**
""".strip()

def generate_customer_from_internal(internal: dict) -> dict:
    """
    Manual-mode helper: creates a simple customer-friendly translation without AI.
    """
    issue = internal.get("primary_issue", "Not specified")
    conditions = internal.get("installation_site_conditions", [])
    concerns = internal.get("potential_concerns", [])
    steps = internal.get("recommended_next_steps", [])

    conditions_line = " ".join(conditions) if conditions else ""
    why = conditions_line if conditions_line else "Further conditions/details were not specified in the notes."

    return {
        "what_we_found": f"We identified a concern related to: {issue}.",
        "why_this_matters": why,
        "what_this_could_lead_to": concerns,
        "recommended_next_steps": steps,
        "priority": internal.get("urgency", "Soon"),
    }

# -------------------- Exports (TXT / DOCX / PDF) -------------------- #

def to_txt_bytes(s: str) -> bytes:
    return (s or "").strip().encode("utf-8")

def md_to_plain_lines(md: str) -> list[str]:
    """
    Very light markdown-to-plain conversion for exports.
    Keeps bullets and headings readable.
    """
    lines = []
    for raw in (md or "").splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        # remove bold markers but keep text
        line = line.replace("**", "")
        lines.append(line)
    return lines

def make_docx_bytes(title: str, md_body: str) -> bytes:
    """
    Creates a simple DOCX (office-friendly).
    Requires python-docx.
    """
    from io import BytesIO
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is not installed. Add it to requirements.txt.") from e

    doc = Document()
    doc.add_heading(title, level=1)

    for line in md_to_plain_lines(md_body):
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def make_pdf_bytes(title: str, md_body: str) -> bytes:
    """
    Creates a simple PDF.
    Requires reportlab.
    """
    from io import BytesIO
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise RuntimeError("reportlab is not installed. Add it to requirements.txt.") from e

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    x = 54
    y = height - 54
    line_height = 14

    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= (line_height * 2)

    c.setFont("Helvetica", 11)
    for line in md_to_plain_lines(md_body):
        # simple page break
        if y <= 54:
            c.showPage()
            y = height - 54
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line[:110])  # simple truncation; keeps PDF generation reliable
        y -= line_height

    c.save()
    return buf.getvalue()

# -------------------- Streamlit UI -------------------- #

st.set_page_config(page_title="Roof Notes → Report", page_icon="🧱", layout="centered")

st.title("🧱 Roof Notes → Dual Report Generator")
st.caption("Creates an INTERNAL manufacturer-safe record (default) + a CUSTOMER-friendly version.")

client = get_client()
if client is None:
    st.warning("OPENAI_API_KEY not found. AI Normalize will not run. Manual Normalize is available below.")
else:
    st.success("AI is available (OPENAI_API_KEY found).")

st.markdown("---")

notes = st.text_area(
    "Paste messy notes / voice transcript",
    placeholder="Example: too cold, glue not sticking, flashing hard to stretch, possible voids at corners.",
    height=180
)

colA, colB, colC = st.columns([1, 1, 1])
with colA:
    normalize_btn = st.button("🧹 Normalize Notes (AI)", type="primary", use_container_width=True)
with colB:
    generate_btn = st.button("📝 Show Reports", use_container_width=True)
with colC:
    clear_btn = st.button("🧽 Clear", use_container_width=True)

if "structured" not in st.session_state:
    st.session_state.structured = None
if "raw_json" not in st.session_state:
    st.session_state.raw_json = None

if clear_btn:
    st.session_state.structured = None
    st.session_state.raw_json = None
    st.rerun()

# ---- AI Normalize ----
if normalize_btn:
    if not notes.strip():
        st.error("Paste some notes first.")
    elif client is None:
        st.error("AI not available. Set OPENAI_API_KEY first (Streamlit Cloud Secrets).")
    else:
        with st.spinner("Normalizing notes…"):
            structured, raw_json, err = normalize_notes_with_gpt(client, notes)

        if err:
            st.error(err)
            st.info("Tip: Use Manual Normalize below to keep working even if AI is unavailable.")
        else:
            st.session_state.structured = structured
            st.session_state.raw_json = raw_json
            st.subheader("Normalized Record (debug)")
            st.json(st.session_state.structured)

            questions = structured.get("clarifying_questions") or []
            if questions:
                st.info("Clarifying questions (answering these will improve accuracy):")
                for q in questions:
                    st.write(f"- {q}")

# ---- Manual Normalize (fallback) ----
st.markdown("---")
st.subheader("Manual Normalize (fallback) — Internal / Manufacturer-Safe (default)")

m_col1, m_col2 = st.columns(2)
with m_col1:
    m_roof_system = st.selectbox(
        "Roof system",
        ["Not specified", "TPO", "EPDM", "PVC", "SBS modified bitumen", "BUR", "Metal", "Shingle"]
    )
    m_primary_issue = st.selectbox(
        "Primary issue",
        ["Not specified", "Active leak", "Ponding", "Open seam / lap", "Flashing concern",
         "Puncture / tear", "Clogged drain / scupper", "Debris", "Moisture concern", "Adhesion / install limitation"]
    )
with m_col2:
    m_location = st.selectbox(
        "Location",
        ["Not specified", "Field", "Perimeter", "At drain / scupper", "At penetration",
         "At parapet / wall detail", "At roof edge / metal", "At corners", "At flashing detail"]
    )
    m_active_leak = st.checkbox("Active leak reported", value=False)

m_service_summary = st.text_area(
    "Service summary (internal)",
    height=90,
    placeholder="Short, manufacturer-safe summary of what happened."
)

m_obs = st.text_area(
    "Observations (one per line)",
    height=110,
    placeholder="- Roof covered in snow\n- Adhesive not bonding\n- Flashing difficult to stretch"
)

m_conditions = st.text_area(
    "Installation / site conditions (one per line)",
    height=90,
    placeholder="- Extreme cold\n- Falling snow\n- Snow-covered substrate"
)

m_concerns = st.text_area(
    "Potential concerns (one per line)",
    height=90,
    placeholder="- Potential voids at corners\n- Risk of incomplete adhesion"
)

m_steps = st.text_area(
    "Recommended next steps (one per line)",
    height=110,
    placeholder="- Reinspect once conditions improve\n- Complete adhesion work under suitable conditions"
)

m_severity = st.selectbox("Severity", ["Low", "Moderate", "High"], index=1)
m_urgency = st.selectbox("Urgency", ["Routine", "Soon", "Immediate"], index=1)

manual_btn = st.button("✅ Use Manual Internal Data", use_container_width=True)

def _lines(txt: str):
    return [x.strip().lstrip("-").strip() for x in (txt or "").splitlines() if x.strip()]

if manual_btn:
    internal = {
        "service_summary": m_service_summary.strip() if m_service_summary.strip() else "Manual entry provided.",
        "roof_system": m_roof_system,
        "primary_issue": m_primary_issue,
        "location": m_location,
        "active_leak_reported": bool(m_active_leak),
        "observations": _lines(m_obs),
        "installation_site_conditions": _lines(m_conditions),
        "potential_concerns": _lines(m_concerns),
        "recommended_next_steps": _lines(m_steps),
        "severity": m_severity,
        "urgency": m_urgency,
    }

    structured = {
        "internal_report": internal,
        "customer_report": generate_customer_from_internal(internal),
        "clarifying_questions": [],
    }

    st.session_state.structured = normalize_and_validate_dual(structured)
    st.session_state.raw_json = None

    st.success("Manual internal record loaded.")
    st.json(st.session_state.structured)

# ---- Show Reports ----
if generate_btn:
    if st.session_state.structured is None:
        st.error("Click 'Normalize Notes (AI)' or 'Use Manual Internal Data' first.")
    else:
        data = st.session_state.structured
        internal = data.get("internal_report", {})
        customer = data.get("customer_report", {})

        internal_md = build_internal_report(internal)
        customer_md = build_customer_report(customer)

        tab1, tab2, tab3 = st.tabs(["Internal (Default)", "Customer-Friendly", "Debug JSON"])

        with tab1:
            st.markdown(internal_md)

        with tab2:
            st.markdown(customer_md)

        with tab3:
            st.json(data)
            if st.session_state.raw_json:
                st.caption("Raw AI output (for debugging):")
                st.code(st.session_state.raw_json, language="json")

        st.markdown("---")
        st.subheader("⬇️ Export (office-friendly)")

        ts = datetime.now().strftime("%Y-%m-%d_%H%M")
        base = f"AllianceRoofing_RoofReport_{ts}"

        col1, col2, col3 = st.columns(3)

        # TXT
        with col1:
            st.download_button(
                "Download Internal (.txt)",
                data=to_txt_bytes(internal_md),
                file_name=f"{base}_INTERNAL.txt",
                mime="text/plain",
                use_container_width=True,
            )
            st.download_button(
                "Download Customer (.txt)",
                data=to_txt_bytes(customer_md),
                file_name=f"{base}_CUSTOMER.txt",
                mime="text/plain",
                use_container_width=True,
            )
            st.download_button(
                "Download Data (.json)",
                data=json.dumps(data, indent=2).encode("utf-8"),
                file_name=f"{base}_DATA.json",
                mime="application/json",
                use_container_width=True,
            )

        # DOCX
        with col2:
            try:
                st.download_button(
                    "Download Internal (.docx)",
                    data=make_docx_bytes("INTERNAL / MANUFACTURER-SAFE REPORT", internal_md),
                    file_name=f"{base}_INTERNAL.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.download_button(
                    "Download Customer (.docx)",
                    data=make_docx_bytes("CUSTOMER-FRIENDLY REPORT", customer_md),
                    file_name=f"{base}_CUSTOMER.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"DOCX export unavailable: {e}")

        # PDF
        with col3:
            try:
                st.download_button(
                    "Download Internal (.pdf)",
                    data=make_pdf_bytes("INTERNAL / MANUFACTURER-SAFE REPORT", internal_md),
                    file_name=f"{base}_INTERNAL.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
                st.download_button(
                    "Download Customer (.pdf)",
                    data=make_pdf_bytes("CUSTOMER-FRIENDLY REPORT", customer_md),
                    file_name=f"{base}_CUSTOMER.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"PDF export unavailable: {e}")

        st.caption(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}  \n"
            "Note: This tool supports documentation. Final assessment should be confirmed by a qualified roofing professional."
        )
